import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm
import os
import math
import pandas as pd
import numpy as np
import re
from datetime import date
import streamlit as st
import json
import glob
from PIL import Image
import smtplib
import docx2pdf 
import shutil
import zipfile
from datetime import datetime
import platform
import matplotlib.pyplot as plt

def User_validation():
    
    
    
        
    f=open("Validation/Validation.json","r")
    past=json.loads(f.read())
    f.close()
    
    now=datetime.now()
    dt_string = now.strftime("%d/%m/%Y %H:%M")
    time_past=datetime.strptime(past['Acceso']["Hora"], "%d/%m/%Y %H:%M")
    
    timesince = now - time_past
    Time_min= int(timesince.total_seconds() / 60)
    
    
   
    bool_negate = Time_min<120
  
    if not bool_negate:
        past['Acceso'].update({"Estado":"Negado"})
        str_json_0=json.dumps(past, indent=4)
        J_f=open("Validation/Validation.json","w")
        J_f.write(str_json_0)
        J_f.close()
    
    

    
    
    
    
    
    bool_aprove= past['Acceso']["Estado"]=="Aprovado"
    
    if not bool_aprove:
        
        colums= st.columns([1,2,1])
        
        with colums[1]:
            #st.image("Imagenes/Escudo_unal.png")
            st.subheader("Ingrese el usuario y contraseña")
            Usuario=st.text_input("Usuario")
            Clave=st.text_input("Contraseña",type="password")
            
        Users=["Gestor Comercial"]
        bool_user = Usuario in Users
        bool_clave = (Clave)==("1234")
    

        bool_user_email = past['Acceso']["User"] == Usuario
    
        bool_time2 = Time_min<1000
        
        bool_1 =  bool_time2 and bool_user_email
        
        
        bool_2 = bool_user and bool_clave
        
        
        if not bool_user_email and  bool_2:
            
            past['Acceso'].update({"User":Usuario,"Estado":"Aprovado","Hora":dt_string})
            str_json_0=json.dumps(past, indent=4)
            J_f=open("Validation/Validation.json","w")
            J_f.write(str_json_0)
            J_f.close()
            
            
        
        
        if not bool_2:
            if (Usuario != "") and (Clave!=""):
                with colums[1]:
                    st.warning("Usuario o contraseña incorrectos.\n\n Por favor intente nuevamente.")
            
        elif bool_2 and not bool_1:
            
            past['Acceso'].update({"User":Usuario,"Estado":"Aprovado","Hora":dt_string})
            str_json_0=json.dumps(past, indent=4)
            J_f=open("Validation/Validation.json","w")
            J_f.write(str_json_0)
            J_f.close()
            
    
            
    
            EMAIL_ADDRESS = 'ASPproyectos2021@gmail.com'
            EMAIL_PASSWORD = 'Samuelmamahuevo2021'
            
            
        
            try:
                
                with smtplib.SMTP('smtp.gmail.com', 587) as smtp:
                    smtp.ehlo()
                    smtp.starttls()
                    smtp.ehlo()
                    
                    smtp.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
                    
                    
                    
                    subject = 'Acceso aplicacion Julia' 
                    body = 'Acceso usuario ' + Usuario +' el '+dt_string 
                    
                    msg = f'Subject: {subject}\n\n{body}'
                    smtp.sendmail(EMAIL_ADDRESS, EMAIL_ADDRESS, msg)
            except:
                pass
            with colums[1]:   
                st.button("Acceder a la aplicación")
        elif bool_2:
            
            past['Acceso'].update({"Estado":"Aprovado","Hora":dt_string,"User":Usuario})
            str_json_0=json.dumps(past, indent=4)
            J_f=open("Validation/Validation.json","w")
            J_f.write(str_json_0)
            J_f.close()
            
            

            
            with colums[1]:   
                st.button("Acceder a la aplicación")
                
    return bool_aprove
def Num_dias(leng):
    if leng==1:
        return "1 día"
    else:
        return str(leng) + " días"
def day_week(dia):
    if dia ==0:
        Dia="Lunes"
    elif dia ==1:
        Dia="Martes"
    elif dia ==2:
        Dia="Miércoles"
    elif dia ==3:
        Dia="Jueves"
    elif dia ==4:
        Dia="Viernes"
    elif dia ==5:
        Dia="Sábado"
    elif dia ==6:
        Dia="Domingo-Festivo"
        
    return Dia
def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)        
 
def Range_fecha(dates):
   
    if len(dates)==1:
        return pd.to_datetime(dates[0]).strftime('%Y-%m-%d')
    else:
        return pd.to_datetime(dates[0]).strftime('%Y-%m-%d')+" hasta "+ pd.to_datetime(dates[-1]).strftime('%Y-%m-%d')

def any2str(obj):
    if isinstance(obj, str):
        return obj
    elif isinstance(obj, int):
        return str(obj)
    elif isinstance(obj, float):
        return str(obj)
    elif math.isnan(obj):
        return ""
    
    
def dt_fechas(data,data_user,Fechas,tipo_dia):
    dt_Final=pd.DataFrame(columns=["Dia","Fecha","Requerimiento","Respaldo"])
    for dia in Fechas:
        data_fecha=data_user[data_user["Fecha"]== dia]
        data_dia_todos=data[data["Fecha"]==dia]
        try:
            d_week=tipo_dia[Tipo_dia["FECHA"]==dia]["TIPO D"].to_numpy()[0]
        except:
            st.warning("Actualizar el calendario del excel extra")
            d_week=day_week(pd.Series(data=dia).dt.dayofweek.to_numpy()[0])
        
        df=pd.DataFrame([[d_week,dia,data_dia_todos["CANTIDAD"].sum(),data_fecha["CANTIDAD"].sum()]],columns=["Dia","Fecha","Requerimiento","Respaldo"])
        dt_Final=dt_Final.append(df, ignore_index=True)
    
    return dt_Final
    
def dt_fechas_2(data,data_user,Fechas,tipo_dia):
    dt_Final=pd.DataFrame(columns=["Dia","Fecha","Requerimiento","Respaldo"])
    for dia in Fechas:
        data_fecha=data_user[data_user["FECHA"]== dia]
        data_dia_todos=data[data["FECHA"]==dia]
        try:
            d_week=tipo_dia[Tipo_dia["FECHA"]==dia]["TIPO D"].to_numpy()[0]
        except:
            st.warning("Actualizar el calendario del excel extra")
            d_week=day_week(pd.Series(data=dia).dt.dayofweek.to_numpy()[0])
        
        df=pd.DataFrame([[d_week,dia,data_dia_todos["CANTIDAD"].sum(),data_fecha["CANTIDAD"].sum()]],columns=["Dia","Fecha","Requerimiento","Respaldo"])
        dt_Final=dt_Final.append(df, ignore_index=True)
    
    return dt_Final

def Mes_espa(mes):
    
    if mes =="01":
        Mes="Enero"
    elif mes =="02":
        Mes="Febrero"
    elif mes =="03":
        Mes="Marzo"
    elif mes =="04":
        Mes="Abril"
    elif mes =="05":
        Mes="Mayo"
    elif mes =="06":
        Mes="Junio"
    elif mes =="07":
        Mes="Julio"
    elif mes =="08":
        Mes="Agosto"
    elif mes =="09":
        Mes="Septiembre"
    elif mes =="10":
        Mes="Octubre"
    elif mes =="11":
        Mes="Noviembre"
    elif mes =="12":
        Mes="Diciembre"
        
    return Mes    

def F_Liq_pag(mes,ano):
    
    if mes%12 ==1:
        Fecha ="Enero"
        
    elif mes%12 ==2:
        Fecha ="Febrero"
        
    elif mes%12 ==3:
        Fecha ="Marzo"
        
    elif mes%12 ==4:
        Fecha ="Abril"
        
    elif mes%12 ==5:
        Fecha ="Mayo"
        
    elif mes%12 ==6:
        Fecha ="Junio"
        
    elif mes%12 ==7:
        Fecha ="Julio"
        
    elif mes%12 ==8:
        Fecha="Agosto"
        
    elif mes%12 ==9:
        Fecha="Septiembre"
        
    elif mes%12 ==10:
        Fecha="Octubre"
        
    elif mes%12 ==11:
        Fecha="Noviembre"
        
    elif mes%12 ==0:
        Fecha="Diciembre"
    
    if mes > 12:
        Fecha += " "+ str(ano+1)
    else:
        Fecha += " "+ str(ano)
    return Fecha       
    
def num2money(num):
    if num < 1e3:
        return str(round(num,2))
    elif num < 1e6:
        return str(round(num*1e3/1e6,2))+ " miles."
    elif num < 1e9:
        return str(round(num*1e3/1e9,2))+ " mill."
    elif num < 1e12:
        return str(round(num*1e3/1e12,2))+ " mil mill."
    
def mes_espa(mes):
    
    if mes =="01":
        Mes="enero"
    elif mes =="02":
        Mes="febrero"
    elif mes =="03":
        Mes="marzo"
    elif mes =="04":
        Mes="abril"
    elif mes =="05":
        Mes="mayo"
    elif mes =="06":
        Mes="junio"
    elif mes =="07":
        Mes="julio"
    elif mes =="08":
        Mes="agosto"
    elif mes =="09":
        Mes="septiembre"
    elif mes =="10":
        Mes="octubre"
    elif mes =="11":
        Mes="noviembre"
    elif mes =="12":
        Mes="diciembre"
    
    
    return Mes
    
def mes_num(mes):
    Opciones2=("Enero","Febrero","Marzo","Abril","Mayo","Junio","Julio","Agosto","Septiembre","Octubre","Noviembre","Diciembre")
    if mes == Opciones2[0]:
        Mes="01"
    elif mes == Opciones2[1]:
        Mes="02"
    elif mes == Opciones2[2]:
        Mes="03"
    elif mes == Opciones2[3]:
        Mes="04"
    elif mes == Opciones2[4]:
        Mes="05"
    elif mes == Opciones2[5]:
        Mes="06"
    elif mes == Opciones2[6]:
        Mes="07"
    elif mes == Opciones2[7]:
        Mes="08"
    elif mes == Opciones2[8]:
        Mes="09"
    elif mes == Opciones2[9]:
        Mes="10"
    elif mes == Opciones2[10]:
        Mes="11"
    elif mes == Opciones2[11]:
        Mes="12"
    
    
    return Mes

def dia_esp(dia):
    
    if dia =="01":
        Dia="1"
    elif dia =="02":
        Dia="2"
    elif dia =="03":
        Dia="3"
    elif dia =="04":
        Dia="4"
    elif dia =="05":
        Dia="5"
    elif dia =="06":
        Dia="6"
    elif dia =="07":
        Dia="7"
    elif dia =="08":
        Dia="8"
    elif dia =="09":
        Dia="9"
    else :
        Dia = dia
 
    
    
    return Dia


def set_font(rows,fila,col,size):

    run=rows[fila].cells[col].paragraphs[0].runs
    font = run[0].font
    font.size= Pt(size)
    font.name = 'Tahoma'
def replace_text_for_image(paragraph, key, value,wid,hei):
    
                            
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, "")
                
        
        for val in value:
            r = paragraph.add_run()
            r.add_picture(val,width=Cm(wid), height=Cm(hei))

        
        


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)

st.set_page_config(
	layout="centered",  # Can be "centered" or "wide". In the future also "dashboard", etc.
	initial_sidebar_state="auto",  # Can be "auto", "expanded", "collapsed"
	page_title="JULIA RD",  # String or None. Strings get appended with "• Streamlit". 
	page_icon="📊",  # String, anything supported by st.image, or None.
)
if User_validation():
    
    
    Opciones1=("Oferta Firme de Respaldo","Certificado de Reintegros","Proyecto 3")
    eleccion=st.sidebar.selectbox('Seleccione el proyecto',Opciones1)
    
    # if False:
    if eleccion==Opciones1[0]:
    
        st.header("Creación ofertas firmes de respaldo")
        st.subheader("Introducción de los documentos")
        
        
        colums= st.columns([1,1,1])
        with colums[0]:                
            uploaded_file_1 = st.file_uploader("Suba el consolidado base")
        with colums[1]:                
            uploaded_file_2 = st.file_uploader("Suba la plantilla del documento")
        with colums[2]:                
            uploaded_file_3 = st.file_uploader("Suba el excel adicional")
        if (uploaded_file_1 is not None) and (uploaded_file_2 is not None) and (uploaded_file_3 is not None):
            try:    
                data=pd.read_excel(uploaded_file_1)
                Extras=pd.read_excel(uploaded_file_3,sheet_name="Usuarios")
                Tipo_dia=pd.read_excel(uploaded_file_3,sheet_name="Calendario")
            except:
                st.warning("Recuerde que el formato del Excel tiene que ser xls")
            data["Fecha"]=data["FECHAINI"].dt.to_pydatetime()
            
            
            
            if data["USUARIO"].isnull().values.any():
                st.warning("Revisar archivo de consolidado base, usuario no encontrado.")   
                data.dropna(subset = ["USUARIO"], inplace=True)
                Users=pd.unique(data["USUARIO"])
            else:
                Users=pd.unique(data["USUARIO"])
        
            
            Extras=pd.read_excel(uploaded_file_3,sheet_name="Usuarios")
            Tipo_dia=pd.read_excel(uploaded_file_3,sheet_name="Calendario")
            template_file_path = uploaded_file_2
            
            today =  date.today()
            fecha=dia_esp(today.strftime("%d")) +" de "+ mes_espa(today.strftime("%m")) +" de "+ today.strftime("%Y")
            
        
            colums= st.columns([1,4,1])
            with colums[1]:
                
                st.subheader("Introducción de las variables")
                
                P_bolsa=st.text_input("Introduzca el Precio de Escasez de Activación",value="10.00")
                P_contrato=st.text_input("Introduzca el precio del contrato [USD]",value="10.00")
                P_TMR=st.text_input("Introduzca el valor de la TRM",value="3,950.00")
                F_TRM = st.date_input("Seleccione la fecha del valor de la TRM:",value=today).strftime("%Y-%m-%d")
                Agente_extra = st.text_input("Introduzca el nombre particular del agente")
                
            columns_2 = st.columns([1,2,2,1])
            
            Opciones2=("Enero","Febrero","Marzo","Abril","Mayo","Junio","Julio","Agosto","Septiembre","Octubre","Noviembre","Diciembre")
            Opciones3=("I","II","III","IV","V")
            
            with columns_2[1]:
                eleccion2=st.selectbox('Seleccione el mes de la OFR',Opciones2)
            with columns_2[2]:
                eleccion3=st.selectbox('Selecciona la semana de la OFR',Opciones3)
            
            
            if Agente_extra:
                Agente_extra="-"+Agente_extra
            else:
                Agente_extra=""
            columns_3 = st.columns([2,1,2])

            with columns_3[1]:
                if platform.system()=='Windows':
                    b=st.checkbox("PDF")
                else:
                    
                    b=False
                a=st.button("Crear los documentos")
                
            Ruta="Documentos/OFR/"+str(today.year)+"/"+mes_num(eleccion2)+"-"+eleccion2 +"/"+ eleccion3
            Ruta_x="Documentos_exportar/"
            
            if os.path.exists(Ruta_x):
            
                shutil.rmtree(Ruta_x)
                Ruta_x=Ruta_x+"/"
            Ruta_x=Ruta_x+"/"
            os.makedirs(Ruta_x, exist_ok=True)
            
        
            if a:
                
                try:
                    path1 = os.path.join(Ruta)
                    shutil.rmtree(path1)
                    os.makedirs(Ruta, exist_ok=True)
                except:
                    os.makedirs(Ruta, exist_ok=True)
                
                
                Ruta_word=Ruta+"/Word"
                Ruta_pdf=Ruta+"/PDF"    
                
                Info ={"Ruta": Ruta, 
                      "File_names": None
                    } 
                   
                File_names=[]
                           
                os.makedirs(Ruta_word, exist_ok=True)
                if b:
                    
                    os.makedirs(Ruta_pdf, exist_ok=True)
                
                zf = zipfile.ZipFile(
                    "Resultado.zip", "w", zipfile.ZIP_DEFLATED)
                my_bar=st.progress(0)
                steps=len(Users)
                steps_done=0
            
                for usuario in Users:
                    
                    data_user=data.copy()
                    data_user=data_user[data_user["USUARIO"]==usuario] 
                    Empresas = pd.unique(data_user["agente1"])
                    
                    Respaldo = data[data["USUARIO"]== usuario]["CANTIDAD"].sum()
                    Fechas = pd.unique(data_user["Fecha"])
                
                    R_fechas = Range_fecha(Fechas)
                    
                    Data_frame_fechas=dt_fechas(data.copy(),data_user,Fechas,Tipo_dia)
                    
                    
                    try:
                        Email = str(Extras[Extras["USUARIO"] == usuario]["CORREO"].values)
                        Porc_come = Extras[Extras["USUARIO"] == usuario]["MARGEN"].values[0]
                    except:

                        Email = ""
                        Porc_come = 0.1
                        st.warning("No hay coincidencia en el Excel de usuarios para: "+usuario)   
                        
                    Email = re.sub("\[|\]|\'|0","",Email)
                    
                    
                    tx_empresas=""
                    for idx ,val in enumerate(Empresas):
                        
                        if len(Empresas)<4:
                            val_2=val[0:3]
                            tx_empresas += val_2
                            if idx==len(Empresas)-1:
                                pass
                            else:
                                tx_empresas +=", "
                            
                        else:
                            tx_empresas += "Los Generadores"
                        
                    P_kwh=float(re.sub(",","",P_TMR))*float(P_contrato)/1000
                    Ingreso=int(P_kwh*Respaldo)
                    C_comer=int(Ingreso*Porc_come)
                    C_GMS=int(Ingreso*4/1000)
                    I_NETO=Ingreso-C_comer-C_GMS
                    
                    if len(Data_frame_fechas.index.values)>13:
                        Enter="\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n"
                    else:
                        Enter=""
                        
                    variables = {
                        "${FECHA}": fecha,
                        "${MES}": eleccion2,
                        "${AGENTES}": tx_empresas,
                        "${USUARIO}": usuario,
                        "${PRECIO_BOLSA}": P_bolsa,
                        "${PRECIO_CONTRATO}": P_contrato,
                        "${FECHA_TRM}": F_TRM,
                        "${PRECIO_TRM}": P_TMR,
                        "${EMAIL_USUARIO}": Email,
                        "${PRECIO_PKWH}":str(round(P_kwh,2)),
                        "${PORC_COMER}":str(int(Porc_come*100))+"%",
                        "${RESPALDO_TOT}":f'{Respaldo:,}',
                        "${INGRESO}":f'{Ingreso:,}',
                        "${COST_COME}":f'{C_comer:,}',
                        "${COST_GMS}":f'{C_GMS:,}',
                        "${INGRESO_NETO}": f'{I_NETO:,}',
                        "${NUM_DIAS}":Num_dias(len(Fechas)),
                        "${RANGO_FECHAS_1}": R_fechas,
                        "${ENTER}": Enter,
                        "${MES_LIQUIDACION}": F_Liq_pag(Opciones2.index(eleccion2)+2,int(today.strftime("%Y"))),
                        "${MES_PAGO}":  F_Liq_pag(Opciones2.index(eleccion2)+3,int(today.strftime("%Y"))),
                        "${INDICADOR}": eleccion3
                    }
                    
                    template_document = docx.Document(template_file_path)
                    
                    for variable_key, variable_value in variables.items():
    
                        for section in template_document.sections:
                            for paragraph in section.header.paragraphs:
                                replace_text_in_paragraph(paragraph, variable_key, variable_value)
    
                        for paragraph in template_document.paragraphs:
                    
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)
                
                        for table in template_document.tables:
                            for col in table.columns:
                                for cell in col.cells:
                                    for paragraph in cell.paragraphs:
                                        replace_text_in_paragraph(paragraph, variable_key, variable_value)
                    
                    rows = template_document.tables[1].rows
                    index_1=Data_frame_fechas.index.values
                    Acum_Req=0
                    Acum_Res=0
                    for idx in index_1:
    
                        
                        rows[int(idx)+1].cells[0].text = Data_frame_fechas.iloc[idx]["Dia"]
                        
                        rows[int(idx)+1].cells[1].text = Data_frame_fechas.iloc[idx]["Fecha"].strftime('%Y-%m-%d')
                        
                        rows[int(idx)+1].cells[2].text = f'{Data_frame_fechas.iloc[idx]["Requerimiento"]:,}'
                        Acum_Req += Data_frame_fechas.iloc[idx]["Requerimiento"]
                        rows[int(idx)+1].cells[3].text = f'{Data_frame_fechas.iloc[idx]["Respaldo"]:,}'
                        Acum_Res += Data_frame_fechas.iloc[idx]["Respaldo"]
                        
                        for idx_2 in range(0,4):
                            run=rows[int(idx)+1].cells[idx_2].paragraphs[0].runs
                            font = run[0].font
                            font.size= Pt(10)
                            font.name = 'Tahoma'
                        
                    for idx in np.arange(len(index_1)+1,37):
                        
                        remove_row(template_document.tables[1], rows[len(index_1)+1])    
                        
                    rows[-1].cells[1].text = Num_dias(len(Fechas))
                    rows[-1].cells[2].text = f'{Acum_Req:,}'
                    rows[-1].cells[3].text = f'{Acum_Res:,}'
                    
                    version=1
                
                    template_document.save(Ruta_x+usuario+"_OFR"+Agente_extra+".docx")
                    zf.write(Ruta_x+usuario+"_OFR"+Agente_extra+".docx")
                    if b:
                        
                        docx2pdf.convert(Ruta_x+"_OFR"+".docx", Ruta_pdf+"/"+usuario+"_OFR"+Agente_extra+".pdf")
                        zf.write(Ruta_x+usuario+"_OFR"+Agente_extra+".pdf")
                    File_names.extend([usuario+"_OFR"+Agente_extra+".docx"])
                    
                    steps_done += 1    
                    my_bar.progress(int(steps_done*100/steps))
                        
                        
                Info.update({"File_names":File_names})
                json_info = json.dumps(Info, indent = 4)
                with open(Ruta_x+'/00_data.json', 'w') as f:
                    json.dump(json_info, f)
                zf.write(Ruta_x+'/00_data.json')    
                zf.close()
                
                
                with open("Resultado.zip", "rb") as fp:
                    with columns_3[1]:
                        btn = st.download_button(
                            label="Descargar resultados",
                            data=fp,
                            file_name="Resultado.zip",
                            mime="application/zip"
                        )
                    
        else:
            st.warning("Necesita subir los tres archivos")   
    #elif False:
    elif eleccion==Opciones1[1]:
        st.header("Creación certificados de reintegros")
        st.subheader("Introducción de los documentos")
        if True:
            colums= st.columns([1,1,1])
            with colums[0]:                
                uploaded_file_1 = st.file_uploader("Suba el documento de liquidación")
            with colums[1]:                
                uploaded_file_2 = st.file_uploader("Suba la plantilla del documento")
            with colums[2]:                
                uploaded_file_3 = st.file_uploader("Suba el excel adicional")
            
                
        else:
            uploaded_file_1="Liquidacion_base.xlsm"
            uploaded_file_2="Certificado_base.docx"
            uploaded_file_3="Excel_extra_certificados.xls"
            
        if (uploaded_file_1 is not None) and (uploaded_file_2 is not None) and (uploaded_file_3 is not None):
            try:
                
                data=pd.read_excel(uploaded_file_1)
                Extras=pd.read_excel(uploaded_file_3,sheet_name="Usuarios")
                Tipo_dia=pd.read_excel(uploaded_file_3,sheet_name="Calendario")
                Agentes=pd.read_excel(uploaded_file_3,sheet_name="Agentes")
            except:
                st.warning("Recuerde que el formato del Excel tiene que ser xls")
            
            data["FECHA"]=data["FECHA"].dt.to_pydatetime()
            
            
            if data["USUARIO"].isnull().values.any():
                st.warning("Revisar archivo de consolidado base, usuario no encontrado.")   
                data.dropna(subset = ["USUARIO"], inplace=True)
                Users=pd.unique(data["USUARIO"])
            else:
                Users=pd.unique(data["USUARIO"])
            
            
            template_file_path = uploaded_file_2
            
            today =  date.today()
            fecha=dia_esp(today.strftime("%d")) +" de "+ mes_espa(today.strftime("%m")) +" de "+ today.strftime("%Y")
            
            
            colums= st.columns([1,4,1])
            with colums[1]:
                
                st.subheader("Introducción de las variables")
                
                F_TRM = st.date_input("Seleccione la fecha del valor de la TRM:",value=today).strftime("%Y-%m-%d")
                
            P_TMR=str(round(data["TRM"].mean(),2))
            columns_2 = st.columns([1,2,2,1])
            
            Opciones2=("Enero","Febrero","Marzo","Abril","Mayo","Junio","Julio","Agosto","Septiembre","Octubre","Noviembre","Diciembre")
            
            with columns_2[1]:
                eleccion3=st.number_input('Seleccione el año del cerficado',value=today.year)
            with columns_2[2]:
                eleccion2=st.selectbox('Seleccione el mes del cerficado',Opciones2)
            
            
                
            columns_3 = st.columns([2,1,2])

            with columns_3[1]:
                if platform.system()=='Windows':
                    b=st.checkbox("PDF")
                else:
                    
                    b=False
                a=st.button("Crear los documentos")
                
            Ruta="Documentos/Certificados/"+str(eleccion3) +"/"+ mes_num(eleccion2)+"-"+eleccion2 
            Ruta_x="Documentos_exportar"
            if os.path.exists(Ruta_x):
                shutil.rmtree(Ruta_x)
                
                Ruta_x=Ruta_x+"/"
            Ruta_x=Ruta_x+"/"
            os.makedirs(Ruta_x, exist_ok=True)
            
            if a:
                try:
                    path1 = os.path.join(Ruta)
                    shutil.rmtree(path1)
                    os.makedirs(Ruta, exist_ok=True)
                except:
                    os.makedirs(Ruta, exist_ok=True)
                
                
                Ruta_word=Ruta+"/Word"
                Ruta_pdf=Ruta+"/PDF"    
                
                Info ={"Ruta": Ruta, 
                      "File_names": None
                    } 
                   
                File_names=[]
                           
                os.makedirs(Ruta_word, exist_ok=True)
                if b:
                    
                    os.makedirs(Ruta_pdf, exist_ok=True)
                
                    
                os.makedirs(Ruta_word, exist_ok=True)
                if b:
                    
                    os.makedirs(Ruta_pdf, exist_ok=True)
                
                zf = zipfile.ZipFile(
                    "Resultado.zip", "w", zipfile.ZIP_DEFLATED)
                my_bar=st.progress(0)
                steps=len(Users)
                steps_done=0
            
                for usuario in Users:
                    
                    data_user=data.copy()
                    data_user=data_user[data_user["USUARIO"]==usuario] 
                    Empresas = pd.unique(data_user["COMPRADOR"])
                    
                    Respaldo = data[data["USUARIO"]== usuario]["CANTIDAD"].sum()
                    Fechas = pd.unique(data_user["FECHA"])
                
                    R_fechas = Range_fecha(Fechas)
                    
                    Data_frame_fechas=dt_fechas_2(data.copy(),data_user,Fechas,Tipo_dia)
                    
                    
                    try:
                        Email = str(Extras[Extras["USUARIO"] == usuario]["CORREO"].values)
                        Porc_come = Extras[Extras["USUARIO"] == usuario]["MARGEN"].values[0]
                    except:
    
                        Email = ""
                        Porc_come = 0.1
                        st.warning("No hay coincidencia en el Excel de usuarios para: "+usuario)   
                        
                    Email = re.sub("\[|\]|\'|0","",Email)
                    
                    
                    tx_empresas=""
                    for idx ,val in enumerate(Empresas):
                        
                        if len(Empresas)<4:
                            val_2=val[0:3]
                            tx_empresas += val_2
                            if idx==len(Empresas)-1:
                                pass
                            else:
                                tx_empresas +=", "
                            
                        else:
                            tx_empresas += "Los Generadores"
                        
                    
                    
                    
                    if len(Data_frame_fechas.index.values)>13:
                        Enter="\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n"
                    else:
                        Enter=""
                    
                    REQ_MAXIMO=Data_frame_fechas["Respaldo"].max()
                    
                    Valor_total=int(round(data_user["P NETO"].sum()))
                    
                    variables = {
                        "${FECHA}": fecha,
                        "${MES}": eleccion2,
                        "${ANO}": str(eleccion3),
                        "${AGENTES}": tx_empresas,
                        "${USUARIO}": usuario,
                        "${OFERTA_MAX}": f'{REQ_MAXIMO:,}',
                        "${FECHA_TRM}": F_TRM,
                        "${P_TRM}": P_TMR,
                        "${EMAIL_USUARIO}": Email,
                        "${PORC_COMER}":str(int(Porc_come*100))+"%",
                        "${RESPALDO_TOT}":"$ "+f'{Valor_total:,}',
                        "${NUM_DIAS}":Num_dias(len(Fechas)),
                        "${RANGO_FECHAS_1}": R_fechas,
                        "${ENTER}": Enter,
                        "${MES_LIQUIDACION}": F_Liq_pag(Opciones2.index(eleccion2)+2,int(today.strftime("%Y"))),
                        "${MES_PAGO}":  F_Liq_pag(Opciones2.index(eleccion2)+3,int(today.strftime("%Y"))),
                        "${INDICADOR}": eleccion3
                    }
                    
                    template_document = docx.Document(template_file_path)
                    
                    for variable_key, variable_value in variables.items():
    
                        for section in template_document.sections:
                            for paragraph in section.header.paragraphs:
                                
                                replace_text_in_paragraph(paragraph, variable_key, variable_value)
    
                        for paragraph in template_document.paragraphs:
                    
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)
                
                        for table in template_document.tables:
                            for col in table.columns:
                                for cell in col.cells:
                                    for paragraph in cell.paragraphs:
                                        replace_text_in_paragraph(paragraph, variable_key, variable_value)
                    
                    
                    
                    
                     
                    rows = template_document.tables[1].rows
                    index_1=Agentes.index.values
                    Total_ingreso=0
                    Total_respaldo=0
                    
                    dt_participa=pd.DataFrame(columns=["OFR","INGRESO","FILA_1","FILA_2"])
                    
                    for idx in index_1:
                        run=rows[int(idx)+3].cells[0].paragraphs[0]
                        rows[int(idx)+3].cells[0].text = "Precio "+Agentes.iloc[idx]["AGENTE"]
                        rows[int(idx)+3].cells[1].text = f'{Agentes.iloc[idx]["PRECIO"]:,}'
                        rows[int(idx)+3].cells[2].text = f'{round(float(re.sub(",","",P_TMR))*Agentes.iloc[idx]["PRECIO"]/1000,2):,}'
                        
                        Respaldo=data_user[data_user["OFR"]==Agentes.iloc[idx]["OFR"]]["CANTIDAD"].sum()
                        
                        rows[int(idx)+13].cells[0].text = "Respaldo "+Agentes.iloc[idx]["AGENTE"]
                        rows[int(idx)+13].cells[1].text = "kWh"
                        rows[int(idx)+13].cells[2].text = f'{Respaldo:,}'
                        
                        
                        for idx_2 in range(0,3):
                            run=rows[int(idx)+3].cells[idx_2].paragraphs[0].runs
                            font = run[0].font
                            font.size= Pt(10)
                            font.name = 'Tahoma'
                        
                        for idx_2 in range(0,3):
                            run=rows[int(idx)+13].cells[idx_2].paragraphs[0].runs
                            font = run[0].font
                            font.size= Pt(10)
                            font.name = 'Tahoma'
                            if idx_2==0:
                                    
                                font.bold= True
                        
                        Total_respaldo += Respaldo
                        
                        df=pd.DataFrame([[Agentes.iloc[idx]["OFR"],Respaldo,int(idx)+3,int(idx)+13]],columns=["OFR","INGRESO","FILA_1","FILA_2"])
                        dt_participa=dt_participa.append(df, ignore_index=True)
                        
                        Total_ingreso += Respaldo*float(re.sub(",","",P_TMR))*Agentes.iloc[idx]["PRECIO"]/1000
                   
                        
                   
                    
                    filas=[18,20,21,22,24]
                    valores=[Total_respaldo,
                             Total_ingreso,
                             Total_ingreso*Porc_come,
                             Total_ingreso*0.004,
                             Valor_total]
                     
                    for idx, val in enumerate(filas):
                        rows[val].cells[2].text = f'{int(round(valores[idx])):,}'
                        run=rows[val].cells[2].paragraphs[0].runs
                        font = run[0].font
                        font.size= Pt(10)
                        font.name = 'Tahoma'
                  
                    
                    
                    contador1=0   
                    contador2=0  
                    for idx in list(dt_participa.index):
                        if dt_participa.iloc[idx]["INGRESO"]==0:
                            remove_row(template_document.tables[1], rows[dt_participa.iloc[idx]["FILA_1"]+contador1])    
                            contador1 -= 1
                            remove_row(template_document.tables[1], rows[dt_participa.iloc[idx]["FILA_2"]+contador1+contador2])  
                            contador2 -= 1
                
                    
                    
                    
                    
                    
                    
                    rows = template_document.tables[2].rows
                    index_1=Data_frame_fechas.index.values
                    Acum_Req=0
                    Acum_Res=0
                    for idx in index_1:
    
                        rows[int(idx)+1].cells[0].text = Data_frame_fechas.iloc[idx]["Fecha"].strftime('%Y-%m-%d')
                        
                        rows[int(idx)+1].cells[1].text = Data_frame_fechas.iloc[idx]["Dia"]
                        
                        
                        rows[int(idx)+1].cells[2].text = f'{Data_frame_fechas.iloc[idx]["Requerimiento"]:,}'
                        Acum_Req += Data_frame_fechas.iloc[idx]["Requerimiento"]
                        rows[int(idx)+1].cells[3].text = f'{Data_frame_fechas.iloc[idx]["Respaldo"]:,}'
                        Acum_Res += Data_frame_fechas.iloc[idx]["Respaldo"]
                        
                        for idx_2 in range(0,4):
                            run=rows[int(idx)+1].cells[idx_2].paragraphs[0].runs
                            font = run[0].font
                            font.size= Pt(10)
                            font.name = 'Tahoma'
                        
                        
                    for idx in np.arange(len(index_1)+1,37):
                        
                        remove_row(template_document.tables[2], rows[len(index_1)+1])    
                        
                    #rows[-1].cells[1].text = Num_dias(len(Fechas))
                    rows[-1].cells[2].text = f'{Acum_Req:,}'
                    rows[-1].cells[3].text = f'{Acum_Res:,}'
                    
                    for idx_2 in range(1,4):
                        run=rows[-1].cells[idx_2].paragraphs[0].runs
                        font = run[0].font
                        font.size= Pt(10)
                        font.name = 'Tahoma'
                    
                    # version=1
                    name_word="Certificado_Reintegros_"+usuario+"_"+eleccion2+"_"+str(eleccion3)+".docx"
                    name_pdf="Certificado_Reintegros_"+usuario+"_"+eleccion2+"_"+str(eleccion3)+".pdf"
                    template_document.save(Ruta_x+name_word)
                    zf.write(Ruta_x+name_word)
                    if b:
                        
                        docx2pdf.convert(Ruta_x+name_word, Ruta_x+name_pdf)
                        zf.write(Ruta_x+name_pdf)
                    File_names.extend([name_word])
                    
                    steps_done += 1    
                    my_bar.progress(int(steps_done*100/steps))
                        
                        
                Info.update({"File_names":File_names})
                json_info = json.dumps(Info, indent = 4)
                with open(Ruta_x+'/00_data.json', 'w') as f:
                    json.dump(json_info, f)
                zf.write(Ruta_x+'/00_data.json')    
                zf.close()
                
                
                with open("Resultado.zip", "rb") as fp:
                    with columns_3[1]:
                        btn = st.download_button(
                            label="Descargar resultados",
                            data=fp,
                            file_name="Resultado.zip",
                            mime="application/zip"
                        )
                
        else:
            st.warning("Necesita subir los tres archivos")   
    #elif True:
    elif eleccion==Opciones1[2]:
        st.header("Creación consolidados mensuales")
        st.subheader("Introducción de los documentos")
        if True:
            colums= st.columns([1,1])
            with colums[0]:                
                uploaded_file_1 = st.file_uploader("Suba el documento de base principal")
            with colums[1]:                
                uploaded_file_2 = st.file_uploader("Suba la plantilla del documento")
            # with colums[2]:                
            #     uploaded_file_3 = st.file_uploader("Suba el excel adicional")
            
                
        else:
            uploaded_file_1="Excel_base.xlsx"
            uploaded_file_2="Plantilla_base.docx"
            # uploaded_file_3="Excel_extra_certificados.xls"
            
        if (uploaded_file_1 is not None) and (uploaded_file_2 is not None): #and (uploaded_file_3 is not None):
            try:
                
                excel_1=pd.ExcelFile(uploaded_file_1)
                Fronteras=pd.read_excel(excel_1,"FRONTERAS")
                Usuarios=pd.read_excel(excel_1,"USUARIOS")
                Year_list=[s for s in excel_1.sheet_names if "INGRESOS" in s]
                Ingresos=pd.read_excel(excel_1,Year_list)
                #Extras=pd.read_excel(uploaded_file_3,sheet_name="Usuarios")
                #Tipo_dia=pd.read_excel(uploaded_file_3,sheet_name="Calendario")
                #Agentes=pd.read_excel(uploaded_file_3,sheet_name="Agentes")
            except:
                st.warning("Recuerde que el formato del Excel tiene que ser xls")
            
            
            
            if Usuarios["USUARIO"].isnull().values.any():
                st.warning("Revisar archivo de consolidado base, usuario no encontrado.")   
                Usuarios.dropna(subset = ["USUARIO"], inplace=True)
                Users=pd.unique(Usuarios["USUARIO"])
                Users = Users[Users != "JULIA-RD"]
            else:
                Users=pd.unique(Usuarios["USUARIO"])
                Users = Users[Users != "JULIA-RD"]
            
            template_file_path = uploaded_file_2
            
            today =  date.today()
            fecha=dia_esp(today.strftime("%d")) +" de "+ mes_espa(today.strftime("%m")) +" de "+ today.strftime("%Y")
            
            
            colums= st.columns([1,4,1])
            with colums[1]:
                
                st.subheader("Introducción de las variables")
                
            columns_2 = st.columns([1,2,2,1])
            
            Opciones2=("Enero","Febrero","Marzo","Abril","Mayo","Junio","Julio","Agosto","Septiembre","Octubre","Noviembre","Diciembre")
            
            with columns_2[1]:
                eleccion3=st.number_input('Seleccione el año del -',value=today.year)
            with columns_2[2]:
                eleccion2=st.selectbox('Seleccione el mes del -',Opciones2)
            
                
            columns_3 = st.columns([2,1,2])

            with columns_3[1]:
                if platform.system()=='Windows':
                    b=st.checkbox("PDF")
                else:
                    
                    b=False
                a=st.button("Crear los documentos")
                
            Ruta="Documentos/Proyecto_3/"+str(eleccion3) +"/"+ mes_num(eleccion2)+"-"+eleccion2 
            Ruta_x="Documentos_exportar"
            if os.path.exists(Ruta_x):
                shutil.rmtree(Ruta_x)
                
                Ruta_x=Ruta_x+"/"
            else:
                Ruta_x=Ruta_x+"/"
            os.makedirs(Ruta_x, exist_ok=True)
            
            if a:
                try:
                    path1 = os.path.join(Ruta)
                    shutil.rmtree(path1)
                    os.makedirs(Ruta, exist_ok=True)
                except:
                    os.makedirs(Ruta, exist_ok=True)
                
                
                Ruta_word=Ruta+"/Word"
                Ruta_pdf=Ruta+"/PDF"    
                
                Info ={"Ruta": Ruta, 
                      "File_names": None
                    } 
                   
                File_names=[]
                           
                os.makedirs(Ruta_word, exist_ok=True)
                if b:
                    
                    os.makedirs(Ruta_pdf, exist_ok=True)
                
                    
                os.makedirs(Ruta_word, exist_ok=True)
                if b:
                    
                    os.makedirs(Ruta_pdf, exist_ok=True)
                
                zf = zipfile.ZipFile(
                    "Resultado.zip", "w", zipfile.ZIP_DEFLATED)
                my_bar=st.progress(0)
                text_rest = st.empty()
                steps=len(Users)
                steps_done=0
                for usuario in Users:
                    
                    text_rest.text("Progreso: "+str(steps_done)+"/"+str(steps) + " -  Actualmente en el usuario: " +usuario )
            
                    data_user=Usuarios.copy()
                    data_user=data_user[data_user["USUARIO"]==usuario] 
                    
                    Num_fronteras = data_user["NUMERO DE FRONTERAS"].values[0]
                    Ene_agregada =  data_user["ENERGÍA AGREGADA"].values[0]
                    Contrato = data_user["CONTRATO"].values[0]
                    Ven_contrato = data_user["VENCIMIENTO CONTRATO"].values[0]
                    Ejecutivo=data_user["EJECUTIVO DE CUENTA"].values[0]
                    Extra_1=data_user["PUNTO FOCAL"].values[0]
                    try:
                        if math.isnan(Extra_1):
                            Extra_1 = ""
                        elif isinstance(Extra_1, int):
                            Extra_1 = str(Extra_1)
                    except:
                        pass
                    try:
                        if math.isnan(Contrato):
                            Contrato = ""
                    except:
                        pass
                    try:
                        if isinstance(Ven_contrato, datetime):
                            Ven_contrato=Ven_contrato.strftime("%d/%m/%Y")
                        elif isinstance(Ven_contrato, int):
                            Ven_contrato = str(Ven_contrato)
                        elif math.isnan(Ven_contrato):
                            Ven_contrato = ""
                    except:
                        pass
                    try:
                        Punto_focal = str(Usuarios[Usuarios["USUARIO"] == usuario]["PUNTO FOCAL"].values)
                    except:
    
                        Punto_focal  = ""
                        st.warning("No hay coincidencia en el Excel de usuarios para: "+usuario)   
                        
                    Punto_focal  = re.sub("\[|\]|\'|0","",Punto_focal )

   
                   
                    template_document = docx.Document(template_file_path)
        

                                        
                    
                    
                    
                    Anos_dias=[s for s in data_user.columns if "DÍAS CERTIFICADOS " in s]
                    Anos=[]
                    
                    for value in Anos_dias:
                        if data_user[value].values[0] != 0:
                            
                            Anos.append(value.replace('DÍAS CERTIFICADOS ', ''))
                            
                    rows = template_document.tables[0].rows
                    
                    contador=0
                    for idx,val in enumerate(Anos):
                        rows[6+contador].cells[0].text = f'{data_user["INGRESO "+val].values[0]:,}'
                        rows[7+contador].cells[0].text = "INGRESO "+val
                        
                        
                        rows[6+contador].cells[2].text = f'{data_user["DÍAS CERTIFICADOS "+val].values[0]:,}'
                        rows[7+contador].cells[2].text = "DÍAS CERTIFICADOS "+val
                        
                        
                        rows[6+contador].cells[3].text = f'{data_user["ENERGIA RESPALDADA "+val].values[0]:,}'
                        rows[7+contador].cells[3].text = "ENERGIA RESPALDADA "+val
                        for idx_2 in [0,2,3]:
                            set_font(rows,6+contador,idx_2,8)
                            set_font(rows,7+contador,idx_2,7)
                        contador += 3
                    cont_enter=0
                    for idx in np.arange(0,28-(6+contador)):
                        
                        remove_row(template_document.tables[0], rows[-1])    
                        cont_enter += 1
                        
                        
                        
                    
                    if len(Anos) > 2:
                        Enter_final="\n"*cont_enter
                        
                        for paragraph in template_document.paragraphs:
                            replace_text_in_paragraph(paragraph, "${IMGENES_2}", '${IMA_INGRESOS}')
                            if "${IMA_INGRESOS}" in  paragraph.text:
                                run = paragraph.add_run()
                                run.add_break(WD_BREAK.PAGE)
                        
                    elif len(Anos) == 1:
                        Enter_final="\n"*int(cont_enter/4)+'${IMA_INGRESOS}'+"\n"*int(cont_enter/4)
                        for paragraph in template_document.paragraphs:
                            replace_text_in_paragraph(paragraph, "${IMGENES_2}", '')
                    elif len(Anos) == 2:
                        Enter_final='${IMA_INGRESOS}'
                        for paragraph in template_document.paragraphs:
                            replace_text_in_paragraph(paragraph, "${IMGENES_2}", '')
                        
                     
                    variables = {
                        "${USUARIO}": usuario,
                        "${EJECUTIVO}":Ejecutivo,
                        "${NUM_FRONTE}": str(Num_fronteras),
                        "${ENE_AGREGADA}": f'{Ene_agregada:,}',
                        "${CONTRATO}": Contrato,
                        "${VEN_CONTRATO}": Ven_contrato,
                        "${PUNTO_FOC}":Extra_1,
                        "${ENTER}":Enter_final
                    }
                    
                        
                        
                    for variable_key, variable_value in variables.items():
    
                        for section in template_document.sections:
                            for paragraph in section.header.paragraphs:
                                
                                replace_text_in_paragraph(paragraph, variable_key, variable_value)
    
                        for paragraph in template_document.paragraphs:
                    
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)
                
                        for table in template_document.tables:
                            for col in table.columns:
                                for cell in col.cells:
                                    for paragraph in cell.paragraphs:
                                        replace_text_in_paragraph(paragraph, variable_key, variable_value)
                    Ruta_img="Imagenes"
                    try:
                        pathx = os.path.join(Ruta_img)
                        shutil.rmtree(pathx)
                        os.makedirs(Ruta_img, exist_ok=True)
                    except:
                        os.makedirs(Ruta_img, exist_ok=True)
                    Imagenes_name=[]
                    color_b = [(0.94,0.49,0.15,1),(0.9,0.13,0.28,1),(0.61,0.61,0.61,1)]
                    Meses=np.array(["Enero","Febrero","Marzo","Abril","Mayo","Junio","Julio","Agosto","Septiembre","Octubre","Noviembre","Diciembre"])
                    for idx,val in enumerate(Anos):
                        name="INGRESOS "+val+".png"
                        data_graph=Ingresos["INGRESOS "+val][Ingresos["INGRESOS "+val]["USUARIO"]==usuario]
                        data_graph= data_graph.dropna(how='all', axis=1)
                        data_graph['MES'] = Meses[data_graph['FECHA'].dt.month-1]
                        if len(data_graph['MES'].unique()) != len(data_graph['MES']):
                            st.warning("Mes repetido revisar "+ "INGRESOS "+val+ " para el usuario "+ usuario)
                        fig = plt.figure(figsize=(9, 3.38))
                        ax = fig.add_axes([0,0,1,1])
                        Meses_gr = data_graph['MES']
                        Ingresos_gr = data_graph["INGRESOS "+val]
                        ax.bar(Meses_gr,Ingresos_gr,color=color_b[idx % 3])
                        plt.xlabel("")
                        plt.ylabel("Ingresos [COP]")
                        plt.title("INGRESOS "+val)
                        max_val=Ingresos_gr.max()
                        if len(data_graph['MES'])>10:
                            plt.xticks(rotation=90)
                        plt.ylim((0,max_val*1.2))
                        
                        for idx,value in enumerate(list(data_graph.index)):
                                plt.text(x = idx , y = data_graph.loc[value]["INGRESOS "+val] + Ingresos_gr.max()*0.05, s = num2money(data_graph.loc[value]["INGRESOS "+val]), size = 9,ha='center',va='center')
                                
                        plt.gca().axes.get_yaxis().set_visible(False)
                        plt.savefig(Ruta_img+"/"+name,dpi=80, bbox_inches='tight',transparent=True)
                        Imagenes_name.extend([Ruta_img+"/"+name])
                    
                    
                    
                    
                    
                    
                    Var_imagenes = {
                        "${IMA_INGRESOS}": Imagenes_name 
                    }
                    
                    for variable_key, variable_value in Var_imagenes.items():
    
                        
                                
                        for paragraph in template_document.paragraphs:
                    
                            replace_text_for_image(paragraph, variable_key,variable_value,10.77,4.04)
                        
                    
                    Imagenes_name=[]
                    Suma_dia=np.array([])
                    color_b = [(0.94,0.49,0.15,1),(0.9,0.13,0.28,1),(0.61,0.61,0.61,1)]
                    dias=np.array(['1. LUNES', '2. MARTES','3. MIERCOLES', '4. JUEVES', '5. VIERNES', '6. SÁBADO', '7. DOMINGO','PROMEDIO'])
                    
                    name="FRONTERAS.png"
                    data_graph=Fronteras[Fronteras["USUARIO"]==usuario]
                    
                    if data_graph.size !=0:
                        data_graph= data_graph.dropna(how='all', axis=1)
                        for idx_2 in dias:
                            suma=data_graph[idx_2].sum()
                            Suma_dia=np.append([suma],Suma_dia)
                            
                        
                        fig = plt.figure(figsize=(9, 2.8))
                        ax = fig.add_axes([0,0,1,1])
    
                        ax.bar(dias,Suma_dia,color=color_b[0])
                        plt.xlabel("")
                        plt.ylabel("Ingresos [COP]")
                        max_val=Suma_dia.max()
            
                        plt.ylim((0,max_val*1.2))
                        
                        for idx,value in enumerate(Suma_dia):
                                plt.text(x = idx , y = value + max_val*0.05, s = f'{int(value):,}', size = 9,ha='center',va='center')
                                
                        #plt.gca().axes.get_yaxis().set_visible(False)
                        plt.savefig(Ruta_img+"/"+name,dpi=200, bbox_inches='tight',transparent=True)
                        Imagenes_name.extend([Ruta_img+"/"+name])
                        
                        
                        
                        
                        
                        
                        Var_imagenes = {
                            "${IMGENES_3}": Imagenes_name 
                        }
                        
                        for variable_key, variable_value in Var_imagenes.items():
        
                            
                                    
                            for paragraph in template_document.paragraphs:
                        
                                replace_text_for_image(paragraph, variable_key,variable_value,23,7.15)
                    
    
                    
    
    
                        rows = template_document.tables[4].rows
                        
                        contador=0
                        for idx,val in enumerate(data_graph.index):
                            
                            
                            rows[idx+1].cells[0].text = any2str(data_graph["FRONTERA COMERCIAL"].values[idx])
                            rows[idx+1].cells[1].text = any2str(data_graph["FRT DDV"].values[idx])
                            rows[idx+1].cells[2].text = any2str(data_graph["COD SIC"].values[idx])
                            rows[idx+1].cells[3].text = any2str(data_graph["PREDIO"].values[idx])
                            rows[idx+1].cells[4].text = any2str(f'{int(data_graph["PROMEDIO"].values[0]):,}')
                            rows[idx+1].cells[5].text = any2str(data_graph["MARGEN"].values[idx])
                            rows[idx+1].cells[6].text = any2str(data_graph["ÚLTIMA DESCONEXION"].values[idx].strftime("%d/%m/%Y"))
                            rows[idx+1].cells[7].text = any2str(str(data_graph["DÍAS CERTIFICADOS"].values[idx]))
                            rows[idx+1].cells[8].text = any2str(data_graph["ACTUALIZACION"].values[idx].strftime("%d/%m/%Y"))
                            try:
                                rows[idx+1].cells[9].text = any2str(data_graph["PROXIMA PRUEBA DDV"].values[idx].strftime("%d/%m/%Y"))
                            except:
                                
                                rows[idx+1].cells[9].text = str(data_graph["PROXIMA PRUEBA DDV"].values[idx])[:10]
                                
    
                            for idx_2 in range(0,10):
                                set_font(rows,idx+1,idx_2,8)
                            contador+=1
                            
                        for idx in np.arange(0,200-contador):
                            remove_row(template_document.tables[4], rows[-1])   
                        
                    else:
                        st.warning("El usuario "+usuario+" no se encuentra en la hoja de Fronteras")
                        rows = template_document.tables[4].rows
                        for idx in np.arange(0,201):
                            remove_row(template_document.tables[4], rows[-1])
                            
                        for paragraph in template_document.paragraphs:
                            replace_text_in_paragraph(paragraph, "${IMGENES_3}", '')
                        
                    
                     
                        
    
                    name_word="Proyecto_3_"+usuario+"_"+eleccion2+"_"+str(eleccion3)+".docx"
                    name_pdf="Proyecto_3_"+usuario+"_"+eleccion2+"_"+str(eleccion3)+".pdf"
                    template_document.save(Ruta_x+name_word)
                    zf.write(Ruta_x+name_word)
                    if b:
                        
                        docx2pdf.convert(Ruta_x+name_word, Ruta_x+name_pdf)
                        zf.write(Ruta_x+name_pdf)
                    File_names.extend([name_word])
                    
                    steps_done += 1    
                    my_bar.progress(int(steps_done*100/steps))
                    
                    text_rest.text("Progreso: "+str(steps_done)+"/"+str(steps) + " -  Actualmente en el usuario: " +usuario )
                    
                        
                        
                Info.update({"File_names":File_names})
                json_info = json.dumps(Info, indent = 4)
                with open(Ruta_x+'/00_data.json', 'w') as f:
                    json.dump(json_info, f)
                zf.write(Ruta_x+'/00_data.json')    
                zf.close()
                
                
                with open("Resultado.zip", "rb") as fp:
                    with columns_3[1]:
                        btn = st.download_button(
                            label="Descargar resultados",
                            data=fp,
                            file_name="Resultado.zip",
                            mime="application/zip"
                        )
                
        else:
            st.warning("Necesita subir los dos archivos") 
    
        
          
   